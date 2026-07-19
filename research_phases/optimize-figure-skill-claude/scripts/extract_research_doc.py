#!/usr/bin/env python3
"""Extract research-document context for Optimize Figure Skill prompt generation.

The script is intentionally local and deterministic. It reads document text and
returns structured sections that an agent can use to generate a scientific
drawing prompt. It does not call any LLM or image model.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


SECTION_ALIASES = {
    "abstract": [
        "abstract",
        "摘要",
        "summary",
    ],
    "methods": [
        "method",
        "methods",
        "methodology",
        "materials and methods",
        "experimental",
        "experiment",
        "approach",
        "方法",
        "研究方法",
        "实验方法",
        "材料与方法",
    ],
    "results": [
        "result",
        "results",
        "experiments",
        "evaluation",
        "结果",
        "实验结果",
        "评估",
    ],
    "discussion": [
        "discussion",
        "讨论",
    ],
    "conclusion": [
        "conclusion",
        "conclusions",
        "结论",
        "总结",
    ],
    "references": [
        "references",
        "bibliography",
        "参考文献",
    ],
    "captions": [
        "caption",
        "captions",
        "figure caption",
        "figure captions",
        "图注",
    ],
}


HEADING_RE = re.compile(
    r"^\s*(?:#{1,6}\s*)?"
    r"(?:(?:\d+(?:\.\d+)*)\s+)?"
    r"(?P<title>"
    r"abstract|summary|introduction|methodology|methods?|materials and methods|"
    r"experimental|experiments?|approach|results?|evaluation|discussion|"
    r"conclusions?|references|bibliography|captions?|figure captions?|"
    r"摘要|引言|绪论|方法|研究方法|实验方法|材料与方法|结果|实验结果|评估|讨论|结论|总结|参考文献|图注"
    r")\s*$",
    re.IGNORECASE,
)


CAPTION_RE = re.compile(
    r"(?im)^\s*((?:fig(?:ure)?\.?|图|表)\s*\d+[^\n]{0,220}(?:\n(?!\s*(?:fig(?:ure)?\.?|图|表)\s*\d+).{0,220}){0,3})"
)


LATEX_COMMAND_REPLACEMENTS = [
    (re.compile(r"\\section\*?\{([^{}]+)\}"), r"\n# \1\n"),
    (re.compile(r"\\subsection\*?\{([^{}]+)\}"), r"\n## \1\n"),
    (re.compile(r"\\subsubsection\*?\{([^{}]+)\}"), r"\n### \1\n"),
    (re.compile(r"\\title\{([^{}]+)\}"), r"\n# \1\n"),
    (re.compile(r"\\caption\{([^{}]+)\}"), r"\nFigure caption: \1\n"),
    (re.compile(r"\\(?:textbf|emph|textit)\{([^{}]+)\}"), r"\1"),
    (re.compile(r"\\cite[t|p]?\{[^{}]*\}"), ""),
    (re.compile(r"\\ref\{[^{}]*\}"), ""),
]


@dataclass
class ExtractionResult:
    source_path: str
    source_type: str
    title: str = ""
    abstract: str = ""
    methods: str = ""
    results: str = ""
    discussion: str = ""
    conclusion: str = ""
    captions: list[str] = field(default_factory=list)
    candidate_contexts: list[dict[str, str]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    parser_notes: list[str] = field(default_factory=list)
    text_length: int = 0
    parser: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_path": self.source_path,
            "source_type": self.source_type,
            "parser": self.parser,
            "title": self.title,
            "abstract": self.abstract,
            "methods": self.methods,
            "results": self.results,
            "discussion": self.discussion,
            "conclusion": self.conclusion,
            "captions": self.captions,
            "candidate_contexts": self.candidate_contexts,
            "warnings": self.warnings,
            "parser_notes": self.parser_notes,
            "text_length": self.text_length,
        }


def run_command(args: list[str], timeout: int = 90) -> str:
    completed = subprocess.run(
        args,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
        check=False,
    )
    stdout = decode_command_output(completed.stdout)
    stderr = decode_command_output(completed.stderr)
    if completed.returncode != 0:
        raise RuntimeError(stderr.strip() or f"command failed: {args}")
    return stdout


def decode_command_output(value: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "cp936", "latin-1"):
        try:
            return value.decode(encoding)
        except UnicodeDecodeError:
            continue
    return value.decode("utf-8", errors="replace")


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def extract_with_pandoc(path: Path) -> str:
    exe = shutil.which("pandoc")
    if not exe:
        raise RuntimeError("pandoc not found")
    return run_command([exe, str(path), "-t", "markdown", "--wrap=none", "-o", "-"])


def summarize_exception(exc: Exception) -> str:
    message = str(exc).strip()
    return message if message else exc.__class__.__name__


def extract_pdf(path: Path) -> tuple[str, str, list[str]]:
    notes: list[str] = []
    pdftotext = shutil.which("pdftotext")
    if pdftotext:
        try:
            text = run_command([pdftotext, "-layout", str(path), "-"])
            if text.strip():
                return text, "pdftotext", notes
            notes.append("pdftotext returned no extractable text.")
        except Exception as exc:
            notes.append(f"pdftotext failed: {summarize_exception(exc)}")
    else:
        notes.append("pdftotext not found.")

    try:
        import pdfplumber  # type: ignore

        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text(x_tolerance=1, y_tolerance=3) or "")
        text = "\n\n".join(pages)
        if text.strip():
            return text, "pdfplumber", notes
        notes.append("pdfplumber returned no extractable text.")
    except ImportError:
        notes.append("pdfplumber not installed.")
    except Exception as exc:
        notes.append(f"pdfplumber failed: {summarize_exception(exc)}")

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text, "pypdf", notes
        notes.append("pypdf returned no extractable text.")
        return "", "no-text", notes
    except ImportError:
        notes.append("pypdf not installed.")
        return "", "no-text", notes
    except Exception as exc:
        notes.append(f"pypdf failed: {summarize_exception(exc)}")
        raise RuntimeError("could not extract PDF text: " + "; ".join(notes)) from exc


def extract_docx(path: Path) -> tuple[str, str, list[str]]:
    notes: list[str] = []
    try:
        text = extract_with_pandoc(path)
        if text.strip():
            return text, "pandoc", notes
        notes.append("pandoc returned no extractable text.")
    except Exception as exc:
        notes.append(f"pandoc failed: {summarize_exception(exc)}")

    try:
        import mammoth  # type: ignore

        with path.open("rb") as fh:
            result = mammoth.extract_raw_text(fh)
        if result.value.strip():
            return result.value, "mammoth", notes
        notes.append("mammoth returned no extractable text.")
    except ImportError:
        notes.append("mammoth not installed.")
    except Exception as exc:
        notes.append(f"mammoth failed: {summarize_exception(exc)}")

    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        paras = [p.text for p in document.paragraphs if p.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        text = "\n".join(paras + table_text)
        if text.strip():
            return text, "python-docx", notes
        notes.append("python-docx returned no extractable text.")
        return "", "no-text", notes
    except ImportError:
        notes.append("python-docx not installed.")
        raise RuntimeError("could not extract DOCX text: " + "; ".join(notes))
    except Exception as exc:
        notes.append(f"python-docx failed: {summarize_exception(exc)}")
        raise RuntimeError("could not extract DOCX text: " + "; ".join(notes)) from exc


def extract_legacy_doc(path: Path) -> tuple[str, str]:
    try:
        text = extract_with_pandoc(path)
        if text.strip():
            return text, "pandoc"
    except Exception:
        pass

    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        document = None
        try:
            document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            text = document.Content.Text
            if text and text.strip():
                return text, "word-com"

            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = Path(tmpdir) / "converted.docx"
                document.SaveAs(str(docx_path), FileFormat=16)
                return extract_docx(docx_path)[0], "word-com-docx"
        finally:
            if document is not None:
                document.Close(False)
            word.Quit()
    except Exception as exc:
        raise RuntimeError(
            "could not extract legacy .doc text; install Word/pywin32, "
            "convert the file to .docx, or provide exported text"
        ) from exc


def expand_latex_inputs(path: Path, seen: set[Path] | None = None) -> str:
    seen = seen or set()
    path = path.resolve()
    if path in seen:
        return ""
    seen.add(path)
    text = read_text(path)

    def replace_input(match: re.Match[str]) -> str:
        raw = match.group(2).strip()
        candidate = (path.parent / raw)
        if candidate.suffix == "":
            candidate = candidate.with_suffix(".tex")
        if candidate.exists() and candidate.is_file():
            return "\n" + expand_latex_inputs(candidate, seen) + "\n"
        return ""

    return re.sub(r"\\(input|include)\{([^{}]+)\}", replace_input, text)


def latex_to_text(path: Path) -> str:
    text = expand_latex_inputs(path)
    text = re.sub(r"%.*", "", text)
    text = re.sub(r"\\begin\{abstract\}", "\n# Abstract\n", text, flags=re.IGNORECASE)
    text = re.sub(r"\\end\{abstract\}", "\n", text, flags=re.IGNORECASE)
    for pattern, replacement in LATEX_COMMAND_REPLACEMENTS:
        text = pattern.sub(replacement, text)
    text = re.sub(r"\\begin\{[^{}]+\}|\\end\{[^{}]+\}", "\n", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{([^{}]*)\})?", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    return text


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def canonical_section(title: str) -> str:
    normalized = title.strip().lower()
    for key, aliases in SECTION_ALIASES.items():
        if any(normalized == alias or normalized.startswith(alias + " ") for alias in aliases):
            return key
    return normalized


def split_sections(text: str) -> dict[str, str]:
    lines = text.splitlines()
    sections: dict[str, list[str]] = {"preamble": []}
    current = "preamble"

    for line in lines:
        match = HEADING_RE.match(line)
        if match:
            current = canonical_section(match.group("title"))
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)

    return {key: normalize_text("\n".join(value)) for key, value in sections.items() if normalize_text("\n".join(value))}


def trim_at_references(text: str) -> tuple[str, bool]:
    lines = text.splitlines()
    for idx, line in enumerate(lines):
        match = HEADING_RE.match(line)
        if match and canonical_section(match.group("title")) == "references":
            return "\n".join(lines[:idx]), True
    return text, False


def extract_title(text: str, sections: dict[str, str], path: Path) -> str:
    title_match = re.search(r"(?m)^#\s+(.+)$", text)
    if title_match:
        return title_match.group(1).strip()
    preamble = sections.get("preamble", "")
    for line in preamble.splitlines():
        line = line.strip()
        if 8 <= len(line) <= 180 and not HEADING_RE.match(line):
            return line
    return path.stem


def compact(value: str, limit: int = 2200) -> str:
    value = normalize_text(value)
    if len(value) <= limit:
        return value
    return value[:limit].rstrip() + "..."


def collect_captions(text: str) -> list[str]:
    captions = []
    for match in CAPTION_RE.finditer(text):
        caption = compact(match.group(1), 700)
        if caption and caption not in captions:
            captions.append(caption)
    return captions[:12]


def make_candidate_contexts(sections: dict[str, str], captions: list[str]) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    preferred = [
        ("abstract", "Abstract"),
        ("methods", "Methods"),
        ("results", "Results"),
        ("conclusion", "Conclusion"),
    ]
    for key, label in preferred:
        if sections.get(key):
            candidates.append({"name": label, "text": compact(sections[key], 1400)})

    if captions:
        candidates.append({"name": "Figure captions", "text": compact("\n\n".join(captions), 1400)})

    combined = "\n\n".join(sections.values())
    keyword_patterns = [
        ("mechanism/process", r"(?i)(mechanism|pathway|pipeline|framework|architecture|workflow|schematic|机制|通路|流程|架构|框架)"),
        ("experiment/system", r"(?i)(apparatus|setup|system|device|instrument|experiment|实验|装置|系统|设备)"),
        ("comparison/result", r"(?i)(compare|comparison|condition|group|result|performance|对比|比较|工况|结果|性能)"),
    ]
    for name, pattern in keyword_patterns:
        match = re.search(pattern, combined)
        if not match:
            continue
        start = max(0, match.start() - 500)
        end = min(len(combined), match.end() + 900)
        snippet = compact(combined[start:end], 1400)
        if snippet and all(snippet != item["text"] for item in candidates):
            candidates.append({"name": name, "text": snippet})

    return candidates[:6]


def extract_document(path: Path) -> ExtractionResult:
    if path.is_dir():
        raise SystemExit("LaTeX folders are not guessed. Please pass the main .tex file path.")

    suffix = path.suffix.lower()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    parser = ""
    parser_notes: list[str] = []
    if suffix == ".pdf":
        raw_text, parser, parser_notes = extract_pdf(path)
    elif suffix in {".docx"}:
        raw_text, parser, parser_notes = extract_docx(path)
    elif suffix == ".doc":
        raw_text, parser = extract_legacy_doc(path)
    elif suffix in {".tex", ".latex"}:
        raw_text = latex_to_text(path)
        parser = "latex"
    elif suffix in {".md", ".markdown", ".txt"}:
        raw_text = read_text(path)
        parser = "plain-text"
    else:
        raise SystemExit(f"Unsupported file type: {suffix}")

    raw_text = normalize_text(raw_text)
    trimmed_text, removed_refs = trim_at_references(raw_text)
    sections = split_sections(trimmed_text)
    captions = collect_captions(raw_text)

    result = ExtractionResult(
        source_path=str(path.resolve()),
        source_type=suffix.lstrip(".") or "unknown",
        title=compact(extract_title(trimmed_text, sections, path), 300),
        abstract=compact(sections.get("abstract", ""), 2200),
        methods=compact(sections.get("methods", ""), 2600),
        results=compact(sections.get("results", ""), 2200),
        discussion=compact(sections.get("discussion", ""), 1600),
        conclusion=compact(sections.get("conclusion", ""), 1600),
        captions=captions,
        candidate_contexts=make_candidate_contexts(sections, captions),
        parser_notes=parser_notes,
        text_length=len(raw_text),
        parser=parser,
    )

    if removed_refs:
        result.warnings.append("References section was detected and excluded from section extraction.")
    if result.text_length < 800:
        result.warnings.append("Extracted text is short; the document may be scanned, image-only, or sparsely parsed.")
    if suffix == ".pdf" and not result.abstract and not result.methods:
        result.warnings.append("PDF section detection was limited; verify snippets before generating the prompt.")
    if suffix == ".pdf" and not raw_text.strip():
        result.warnings.append("No extractable PDF text was found; this may be a scanned or image-only PDF. Use OCR or paste the abstract/methods text.")
    if suffix == ".docx" and not raw_text.strip():
        result.warnings.append("No extractable DOCX text was found; paste the relevant sections or export the document as plain text.")
    if suffix in {".tex", ".latex"}:
        result.warnings.append("LaTeX macros were simplified; verify equations and specialized notation manually.")

    return result


def to_markdown(result: ExtractionResult) -> str:
    data = result.to_dict()
    lines = [
        f"# {data['title'] or Path(data['source_path']).stem}",
        "",
        f"- Source: `{data['source_path']}`",
        f"- Type: `{data['source_type']}`",
        f"- Parser: `{data['parser']}`",
        f"- Extracted characters: {data['text_length']}",
    ]
    if data["warnings"]:
        lines.append("- Warnings: " + "; ".join(data["warnings"]))
    if data["parser_notes"]:
        lines.append("- Parser notes: " + "; ".join(data["parser_notes"]))
    lines.append("")

    for key, heading in [
        ("abstract", "Abstract"),
        ("methods", "Methods"),
        ("results", "Results"),
        ("discussion", "Discussion"),
        ("conclusion", "Conclusion"),
    ]:
        if data[key]:
            lines.extend([f"## {heading}", data[key], ""])

    if data["captions"]:
        lines.append("## Captions")
        for item in data["captions"]:
            lines.append(f"- {item}")
        lines.append("")

    if data["candidate_contexts"]:
        lines.append("## Candidate Context Snippets")
        for item in data["candidate_contexts"]:
            lines.extend([f"### {item['name']}", item["text"], ""])

    return "\n".join(lines).rstrip() + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract research-document context for Optimize Figure Skill prompt generation.")
    parser.add_argument("path", help="Path to .docx, .doc, .pdf, .tex, .latex, .md, or .txt")
    parser.add_argument("--format", choices=["markdown", "json"], default="markdown", help="Output format")
    args = parser.parse_args()

    result = extract_document(Path(args.path))
    if args.format == "json":
        print(json.dumps(result.to_dict(), ensure_ascii=False, indent=2))
    else:
        print(to_markdown(result), end="")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except RuntimeError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(1)
    except BrokenPipeError:
        raise SystemExit(1)
